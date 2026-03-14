#!/usr/bin/env python3
"""
Robust Local NLP Requirement Extractor (Heuristic + NLP)
--------------------------------------------------------
- Reads DOCX, PDF, XLSX, CSV, and TXT requirement(-ish) documents
- Detects requirement-like sentences even without normative keywords
- Uses multiple signals:
    * keywords (shall/should/may/must/will/needs to/...)
    * pattern-based cues (imperative, functional verbs, SVO structure, passive)
    * section context
    * table context
    * spaCy NLP parsing
- Auto-detects requirement ID formats (no "REQ" assumption)
- Supports multiple ID schemes in one document
- Parses subject/action/object using spaCy
- Exports a JSON database (top-level array) with confidence and metadata

Dependencies:
    pip install python-docx pdfplumber spacy pandas openpyxl
    python -m spacy download en_core_web_sm

Usage:
    python req_extract.py input_file output.json
    # input_file can be .docx, .pdf, .xlsx, .csv, or .txt
    # If no args are provided, the script prompts for the input file path.
"""

import warnings
warnings.filterwarnings("ignore", category=Warning)

import re
import json
import sys
import csv
from pathlib import Path
from typing import List, Dict, Optional

import spacy
import pdfplumber

import pandas as pd
#import pdfplumber
from docx import Document

from file_checks import classifiers_exist
from training import train_classifiers_csv
from evaluation import run_kfold_evaluation
from batch import batch_process_csv
import time

import matplotlib.pyplot as plt



# Load spaCy model once (AI component)
nlp = spacy.load("en_core_web_sm")

# ------------------------------------------------------------
# 1. Generalized, Multi-Scheme ID Detection
# ------------------------------------------------------------


ID_PATTERNS = [
    r"[A-Za-z]+[-_ ]?\d+(\.\d+)*",   # Prefix + number, optional hierarchy
    r"\d+(\.\d+)+",                  # Hierarchical numeric IDs (1.2, 1.2.3)
    r"[A-Za-z]+[-_ ]?\d+",           # Prefix + number (simple)
]

COMPILED_PATTERNS = [re.compile(p) for p in ID_PATTERNS]


def extract_req_id(text: str) -> Optional[str]:
    """Extract the first matching ID from the text using generalized patterns."""
    for pattern in COMPILED_PATTERNS:
        match = pattern.search(text)
        if match:
            return match.group(0)
    return None


# ------------------------------------------------------------
# 2. Document Loaders
# ------------------------------------------------------------

def load_docx(path: Path) -> List[Dict]:
    doc = Document(path)
    blocks = []

    # Add all paragraphs (including headings)
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            blocks.append({"type": "paragraph", "text": text})

    # Add all table cells as paragraphs
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    blocks.append({"type": "paragraph", "text": cell_text})

    return blocks


def load_pdf(path: Path) -> List[Dict]:
    blocks = []
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            text = page.extract_text()
            if not text:
                continue
            for line in text.split("\n"):
                line = line.strip()
                if line:
                    blocks.append({"type": "paragraph", "text": line})
    return blocks



def load_xlsx(path: Path) -> List[Dict]:
    """
    Load xlsx file and create a block for each non-empty cell.
    This ensures that text within a single cell (even multi-line text) 
    stays together and is never split.
    """
    blocks = []
    df = pd.read_excel(path)
    for _, row in df.iterrows():
        for cell in row:
            if pd.notnull(cell):
                text = str(cell).strip()
                if text:
                    blocks.append({"type": "cell", "text": text})
    return blocks

def load_csv(path: Path) -> List[Dict]:
    """
    Load CSV file and create a block for each non-empty cell.
    This ensures that text within a single cell (even multi-line text) 
    stays together and is never split.
    """
    blocks = []
    df = pd.read_csv(path)
    for _, row in df.iterrows():
        for cell in row:
            if pd.notnull(cell):
                text = str(cell).strip()
                if text:
                    blocks.append({"type": "cell", "text": text})
    return blocks

def load_txt(path: Path) -> List[Dict]:
    blocks = []
    with open(path, 'r', encoding='utf-8') as f:
        for line in f:
            text = line.strip()
            if text:
                blocks.append({"type": "line", "text": text})
    return blocks

def load_document(path: Path) -> List[Dict]:
    loaders = {
        ".docx": load_docx,
        ".pdf": load_pdf,
        ".xlsx": load_xlsx,
        ".csv": load_csv,
        ".txt": load_txt,
    }
    ext = path.suffix.lower()
    if ext in loaders:
        return loaders[ext](path)
    raise ValueError(f"Unsupported file format: {path.suffix}")





# ------------------------------------------------------------
# 3. Requirement-Like Detection (Heuristics + NLP)
# ------------------------------------------------------------

# Expanded normative / modal / permissive phrases
SINGLE_WORD_NORMATIVE = {
    "shall", "should", "may", "must", "will", "can"
}

MULTIWORD_NORMATIVE = [
    "needs to",
    "is required to",
    "has to",
    "is able to",
    "is expected to",
    "is permitted to",
    "must not",
    "shall not",
]

# Functional verbs that often indicate behavior
FUNCTIONAL_VERBS = {
    "load", "save", "store", "export", "import", "generate", "create",
    "delete", "update", "edit", "validate", "verify", "authenticate",
    "authorize", "encrypt", "decrypt", "log", "record", "display",
    "show", "render", "calculate", "compute", "send", "receive",
    "transmit", "print", "filter", "search", "sort", "notify",
    "alert", "backup", "restore", "sync", "synchronize"
}

# Section keywords that hint at requirement context
REQUIREMENT_SECTION_HINTS = [
    "requirement", "requirements", "functional requirements",
    "system requirements", "software requirements", "user requirements"
]


def contains_multiword_phrase(text: str, phrases: List[str]) -> bool:
    lowered = text.lower()
    return any(p in lowered for p in phrases)


def contains_singleword(text: str, words: set) -> bool:
    tokens = text.lower().split()
    return any(w in tokens for w in words)


def section_is_requirement_like(section: Optional[str]) -> bool:
    if not section:
        return False
    lowered = section.lower()
    return any(h in lowered for h in REQUIREMENT_SECTION_HINTS)


def analyze_requirement_likeness(
    text: str,
    section: Optional[str] = None,
    in_table: bool = False
) -> Dict:
    """
    Multi-signal requirement detection:
      - normative/modal keywords (single + multiword)
      - functional verbs
      - imperative form
      - subject-verb-object structure
      - passive constructions
      - section context
      - table context
      - spaCy-based NLP analysis

    Returns:
      {
        "is_requirement": bool,
        "confidence": float (0.0 - 1.0),
        "signals": [str, ...]
      }
    """
    signals = []
    score = 0.0

    # Quick lexical checks
    if contains_singleword(text, SINGLE_WORD_NORMATIVE):
        signals.append("normative_singleword")
        score += 0.25

    if contains_multiword_phrase(text, MULTIWORD_NORMATIVE):
        signals.append("normative_multiword")
        score += 0.25

    # NLP analysis
    doc = nlp(text)

    has_verb = any(t.pos_ == "VERB" for t in doc)
    if has_verb:
        signals.append("has_verb")
        score += 0.1

    # Functional verbs
    if any(t.lemma_.lower() in FUNCTIONAL_VERBS for t in doc if t.pos_ == "VERB"):
        signals.append("functional_verb")
        score += 0.15

    # Subject-Verb-Object structure
    has_subject = any(t.dep_ in ("nsubj", "nsubjpass") for t in doc)
    has_object = any(t.dep_ in ("dobj", "pobj", "attr") for t in doc)
    if has_subject and has_verb and has_object:
        signals.append("svo_structure")
        score += 0.15

    # Imperative: first token is a verb, no explicit subject
    if doc and doc[0].pos_ == "VERB" and not has_subject:
        signals.append("imperative_form")
        score += 0.15

    # Passive voice: auxpass or "be" + past participle
    passive = any(t.dep_ == "auxpass" for t in doc)
    if passive:
        signals.append("passive_voice")
        score += 0.05

    # Section context
    if section_is_requirement_like(section):
        signals.append("requirement_section_context")
        score += 0.1

    # Table context
    if in_table:
        signals.append("table_context")
        score += 0.1

    # Normalize score to max 1.0
    score = min(score, 1.0)

    # Decision threshold: tuneable
    is_req = score >= 0.35

    return {
        "is_requirement": is_req,
        "confidence": round(score, 3),
        "signals": signals
    }


# ------------------------------------------------------------
# 4. NLP Parsing (AI Part: Subject/Action/Object)
# ------------------------------------------------------------

def parse_subject_action_object(text: str) -> Dict:
    """
    Uses spaCy's statistical NLP model to extract:
        - subject (nsubj)
        - action (ROOT verb)
        - object (dobj, pobj, attr)
    """
    doc = nlp(text)

    subject = None
    action = None
    obj = None

    for token in doc:
        if token.dep_ in ("nsubj", "nsubjpass") and subject is None:
            subject = token.text
        if token.dep_ == "ROOT" and token.pos_ == "VERB" and action is None:
            action = token.lemma_
        if token.dep_ in ("dobj", "pobj", "attr") and obj is None:
            obj = token.text

    return {
        "subject": subject,
        "action": action,
        "object": obj,
        "issues": []
    }


# ------------------------------------------------------------
# 5. Requirement Extraction
# ------------------------------------------------------------

def extract_requirements(blocks: List[Dict], source_file: str) -> List[Dict]:
    results = []
    current_section = None
    auto_counter = 1

    def next_req_id(text: str) -> str:
        nonlocal auto_counter
        req_id = extract_req_id(text)
        if not req_id:
            req_id = f"AUTO-{auto_counter:04d}"
            auto_counter += 1
        return req_id

    def build_result(
        req_id: str,
        source_section: Optional[str],
        source_type: str,
        raw_text: str,
        analysis: Dict,
        parsed: Dict,
        priority: Optional[str],
    ) -> Dict:
        return {
            "req_id": req_id,
            "source_section": source_section,
            "source_type": source_type,
            "raw_text": raw_text,
            "req_confidence": analysis["confidence"],
            "req_signals": ";".join(analysis["signals"]),
            "subject": parsed["subject"],
            "action": parsed["action"],
            "object": parsed["object"],
            "priority": priority,
            "issues": parsed["issues"],
            "source_file": source_file,
        }

    for block in blocks:

        # Track headings
        if block["type"] == "heading":
            current_section = block["text"]
            continue

        # Paragraph requirements
        if block["type"] == "paragraph":
            text = block["text"]
            analysis = analyze_requirement_likeness(
                text,
                section=current_section,
                in_table=False
            )
            if not analysis["is_requirement"]:
                continue

            req_id = next_req_id(text)
            parsed = parse_subject_action_object(text)
            results.append(
                build_result(
                    req_id,
                    current_section,
                    "paragraph",
                    text,
                    analysis,
                    parsed,
                    None,
                )
            )

        # Table requirements (DOCX tables)
        if block["type"] == "table":
            rows = block["rows"]
            if not rows:
                continue

            header = [h.lower() for h in rows[0]]

            # Detect columns (best-effort)
            try:
                id_idx = header.index("id")
            except ValueError:
                id_idx = None

            text_idx = None
            for candidate in ["requirement", "requirement text", "description", "text"]:
                if candidate in header:
                    text_idx = header.index(candidate)
                    break

            prio_idx = header.index("priority") if "priority" in header else None

            # If a recognized requirement column exists, use it; otherwise, analyze all cells in all rows
            if text_idx is not None:
                for row in rows[1:]:
                    if text_idx >= len(row):
                        continue
                    req_text = row[text_idx].strip()
                    if not req_text:
                        continue

                    analysis = analyze_requirement_likeness(
                        req_text,
                        section=current_section,
                        in_table=True
                    )
                    if not analysis["is_requirement"]:
                        continue

                    # ID extraction: from ID column if present, else from text, else AUTO
                    req_id = None
                    if id_idx is not None and id_idx < len(row):
                        req_id = extract_req_id(row[id_idx])
                    if not req_id:
                        req_id = next_req_id(req_text)

                    parsed = parse_subject_action_object(req_text)

                    priority = row[prio_idx].strip() if prio_idx is not None and prio_idx < len(row) else None
                    results.append(
                        build_result(
                            req_id,
                            current_section,
                            "table",
                            req_text,
                            analysis,
                            parsed,
                            priority,
                        )
                    )
            else:
                # No recognized requirement column: analyze all cells in all rows
                for row in rows[1:]:
                    for cell in row:
                        req_text = str(cell).strip()
                        if not req_text:
                            continue
                        analysis = analyze_requirement_likeness(
                            req_text,
                            section=current_section,
                            in_table=True
                        )
                        if not analysis["is_requirement"]:
                            continue
                        req_id = next_req_id(req_text)
                        parsed = parse_subject_action_object(req_text)
                        results.append(
                            build_result(
                                req_id,
                                current_section,
                                "table",
                                req_text,
                                analysis,
                                parsed,
                                None,
                            )
                        )

        # Cell requirements (XLSX cells)
        if block["type"] == "cell":
            text = block.get("text", "")
            analysis = analyze_requirement_likeness(text)
            if not analysis["is_requirement"]:
                continue
            req_id = next_req_id(text)
            parsed = parse_subject_action_object(text)
            results.append(
                build_result(
                    req_id,
                    current_section,
                    "cell",
                    text,
                    analysis,
                    parsed,
                    None,
                )
            )

        # Row requirements (CSV rows)
        if block["type"] == "row":
            # Try to split row into cells if possible
            text = block.get("text", "")
            # If the row is a string, try splitting by comma or tab
            cells = [text]
            if "," in text:
                cells = [c.strip() for c in text.split(",") if c.strip()]
            elif "\t" in text:
                cells = [c.strip() for c in text.split("\t") if c.strip()]
            for cell in cells:
                analysis = analyze_requirement_likeness(cell)
                if not analysis["is_requirement"]:
                    continue
                req_id = next_req_id(cell)
                parsed = parse_subject_action_object(cell)
                results.append(
                    build_result(
                        req_id,
                        current_section,
                        "row",
                        cell,
                        analysis,
                        parsed,
                        None,
                    )
                )

    return results


# ------------------------------------------------------------
# 6. JSON Export
# ------------------------------------------------------------

EXTRACTOR_VERSION = "Extractor_LocalNLP_Ver2"


def export_json_database(rows: List[Dict], out_path: Path, source_file: str):
    if not rows:
        print("\nNo requirements found.")
        return

    output = []
    for idx, row in enumerate(rows, start=1):
        output.append({
            "id": idx,
            "raw_text": row.get("raw_text", ""),
            "req_confidence": row.get("req_confidence"),
            "subject": row.get("subject"),
            "action": row.get("action"),
            "object": row.get("object"),
            "source": source_file,
            "extractor_version": EXTRACTOR_VERSION,
        })

    with out_path.open("w", encoding="utf-8") as f:
        json.dump(output, f, indent=2)

    print(f"\nExported {len(output)} requirements to {out_path}")


def export_csv_raw_text(rows: List[Dict], out_path: Path, source_file: str):
    """Export only raw text of extracted requirements to CSV."""
    if not rows:
        print("No requirements found.")
        return

    with out_path.open("w", encoding="utf-8", newline="") as f:
        writer = csv.writer(f)
        writer.writerow(["raw_text"])
        for row in rows:
            writer.writerow([row.get("raw_text", "")])

    print(f"\nExported {len(rows)} requirements (raw text) to {out_path}")


def export_csv_database(rows: List[Dict], out_path: Path, source_file: str):
    """Export extracted requirements to CSV with all attributes matching JSON output."""
    if not rows:
        print("No requirements found.")
        return

    fieldnames = ["id", "raw_text", "req_confidence", "subject", "action", "object", "source", "extractor_version"]
    
    with out_path.open("w", encoding="utf-8", newline="") as f:
        writer = csv.DictWriter(f, fieldnames=fieldnames)
        writer.writeheader()
        
        for idx, row in enumerate(rows, start=1):
            writer.writerow({
                "id": idx,
                "raw_text": row.get("raw_text", ""),
                "req_confidence": row.get("req_confidence"),
                "subject": row.get("subject"),
                "action": row.get("action"),
                "object": row.get("object"),
                "source": source_file,
                "extractor_version": EXTRACTOR_VERSION,
            })

    print(f"\nExported {len(rows)} requirements to {out_path}")


# ------------------------------------------------------------
# 7. CLI
# ------------------------------------------------------------

def main():

    print("\nExtractor_LocalNLP_Ver2.py: multi-format input enabled (docx, pdf, xlsx, csv, txt)")
    if len(sys.argv) == 1:
        selected = input("\nInput file path: ").strip()
        if not selected:
            print("No input file selected.")
            sys.exit(1)

        input_path = Path(selected)
        output_name = input_path.stem + "_extracted"
    elif len(sys.argv) == 3:
        input_path = Path(sys.argv[1])
        output_name = Path(sys.argv[2]).stem
    else:
        print("Usage: python req_extract.py input_file output.json\n  input_file can be .docx, .pdf, .xlsx, .csv, or .txt")
        sys.exit(1)

    output_dir = input_path.parent
    output_path = output_dir / ("Input_Restructured.json")


    blocks = load_document(input_path)

    # Step 3: Requirement-likeness analysis for all blocks
    req_likeness_path = output_dir / ("extracted_req_likeness.txt")
    with req_likeness_path.open('w', encoding='utf-8') as f:
        for block in blocks:
            type_ = block.get('type', '')
            text = block.get('text', '')
            # Only analyze text blocks (skip tables/rows with no text)
            if type_ in ("paragraph", "row", "line", "cell") and text:
                analysis = analyze_requirement_likeness(text)
                is_req = analysis["is_requirement"]
                confidence = analysis["confidence"]
                signals = analysis["signals"]
                f.write(f"type: {type_}\n")
                f.write(f"text: {text}\n")
                f.write(f"is_requirement: {is_req}\n")
                f.write(f"confidence: {confidence}\n")
                f.write(f"signals: {signals}\n")
                f.write("---\n")
            elif type_ == "table":
                f.write(f"type: table\nrows: {block.get('rows', '')}\n---\n")
    print(f"\nExported requirement-likeness analysis to {req_likeness_path}")

    rows = extract_requirements(blocks, input_path.name)
    export_json_database(rows, output_path, input_path.name)
    
    # Export CSV with all attributes
    csv_path = output_dir / ("Input_Restructured.csv")
    export_csv_database(rows, csv_path, input_path.name)

if __name__ == "__main__":
    main()

#################################################################################################################################
    #New Script#
    
tic = time.perf_counter()

print("\n=== Hybrid Conformance Checker ===")

if not classifiers_exist():
    print("\nNo trained classifiers found. Training models for the first time...")
    run_kfold_evaluation(k=5)
    train_classifiers_csv()
else:
    print("\nClassifiers found. Skipping training.")

print("\nRunning batch conformance checking...")
batch_process_csv()

print("\nDone.")

toc = time.perf_counter()
print(f"Elapsed: {toc - tic:.6f} seconds")